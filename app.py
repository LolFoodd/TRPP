import os
import sqlite3
from datetime import date, datetime

from docx import Document
from flask import Flask, render_template, redirect, url_for, abort, request, send_file

app = Flask(__name__)


def get_db_connection():
    conn = sqlite3.connect('database.db')
    conn.row_factory = sqlite3.Row
    return conn


def get_contracts():
    conn = get_db_connection()
    contracts_db = conn.execute('SELECT * FROM contracts').fetchall()
    conn.close()
    return contracts_db


def get_contract(contract_id):
    conn = get_db_connection()
    contract_id_db = conn.execute('SELECT * FROM contracts WHERE id = ?',
                                  (contract_id,)).fetchone()
    conn.close()
    if contract_id_db is None:
        abort(404)
    return contract_id_db


def add_new_contract(number, services_type_id, start_price, discount, client_id, employee_id):
    conn = get_db_connection()
    conn.execute(
        "INSERT INTO 'contracts' ('number', 'date', 'services_type_id', 'start_price', 'discount', 'finish_price', 'client_id', 'employee_id')  VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
        (number, date.today(), services_type_id, start_price, discount,
         int(int(start_price) * (1 - int(discount) / 100)),
         client_id, employee_id)).fetchall()
    conn.commit()
    conn.close()


def add_update_contract(id, number, services_type_id, start_price, discount, client_id, employee_id):
    conn = get_db_connection()
    conn.execute(
        "UPDATE contracts SET number = ?, date = ?, services_type_id = ?, start_price = ?, discount = ?, finish_price = ?, client_id = ?, employee_id = ? WHERE id = ?",
        (number, date.today(), services_type_id, start_price, discount,
         int(int(start_price) * (1 - int(discount) / 100)),
         client_id, employee_id, id)).fetchall()
    conn.commit()
    conn.close()


def get_clients():
    conn = get_db_connection()
    clients_db = conn.execute('SELECT * FROM clients').fetchall()
    conn.close()
    return clients_db


def get_client(client_id):
    conn = get_db_connection()
    client_id_db = conn.execute('SELECT * FROM clients WHERE id = ?',
                                (client_id,)).fetchone()
    conn.close()
    if client_id_db is None:
        abort(404)
    return client_id_db


def add_new_client(name, email, phone_number):
    conn = get_db_connection()
    conn.execute("INSERT INTO 'clients' ('name', 'email', 'phone_number')  VALUES (?, ?, ?)",
                 (name, email, phone_number)).fetchall()
    conn.commit()
    conn.close()


def add_update_client(id, name, email, phone_number):
    conn = get_db_connection()
    conn.execute("UPDATE clients SET name = ?, email=?, phone_number=? WHERE id = ?",
                 (name, email, phone_number, id)).fetchall()
    conn.commit()
    conn.close()


def get_employees():
    conn = get_db_connection()
    employees_db = conn.execute('SELECT * FROM employees').fetchall()
    conn.close()
    return employees_db


def get_employee(employee_id):
    conn = get_db_connection()
    employee_id_db = conn.execute('SELECT * FROM employees WHERE id = ?',
                                  (employee_id,)).fetchone()
    conn.close()
    if employee_id_db is None:
        abort(404)
    return employee_id_db


def add_new_employee(name, email, phone_number, position):
    conn = get_db_connection()
    conn.execute("INSERT INTO 'employees' ('name', 'email', 'phone_number', 'position')  VALUES (?, ?, ?, ?)",
                 (name, email, phone_number, position)).fetchall()
    conn.commit()
    conn.close()


def add_update_employee(id, name, email, phone_number, position):
    conn = get_db_connection()
    conn.execute("UPDATE employees SET name = ?, email=?, phone_number=?, position=? WHERE id = ?",
                 (name, email, phone_number, position, id)).fetchall()
    conn.commit()
    conn.close()


def get_type_services():
    conn = get_db_connection()
    type_services_db = conn.execute('SELECT * FROM type_services').fetchall()
    conn.close()
    return type_services_db


def get_type_service(type_service_id):
    conn = get_db_connection()
    type_service_id_db = conn.execute('SELECT * FROM type_services WHERE id = ?',
                                      (type_service_id,)).fetchone()
    conn.close()
    if type_service_id_db is None:
        abort(404)
    return type_service_id_db


def add_new_type_service(service):
    conn = get_db_connection()
    conn.execute("INSERT INTO type_services ('service') VALUES (?)",
                 (service,)).fetchall()
    conn.commit()
    conn.close()


def add_update_type_service(id, service):
    conn = get_db_connection()
    conn.execute("UPDATE type_services SET service = ? WHERE id = ?",
                 (service, id)).fetchall()
    conn.commit()
    conn.close()


def select_for_report(client_id, date):
    conn = get_db_connection()
    contr = conn.execute('SELECT * FROM contracts WHERE client_id=?',
                         (client_id,)).fetchall()
    conn.close()

    date_format = '%Y-%m-%d'
    date1 = datetime.strptime(date, date_format)
    list = []
    for i in contr:
        date2 = datetime.strptime(i["date"], date_format)
        if date2 >= date1:
            list.append(i)
    return list


@app.route('/')
def index():
    return redirect(url_for('contracts'))


@app.route('/contracts')
def contracts():
    contracts_db = get_contracts()
    map_client = {}
    map_employee = {}
    map_type_service = {}
    for contract in contracts_db:
        map_client[contract['client_id']] = get_client(contract['client_id'])['name']
        map_employee[contract['employee_id']] = get_employee(contract['employee_id'])['name']
        map_type_service[contract['services_type_id']] = get_type_service(contract['services_type_id'])['service']
    return render_template('contracts.html', contracts=contracts_db, clients=map_client, employees=map_employee,
                           services=map_type_service)


@app.route('/contract/<int:contract_id>')
def contract(contract_id):
    contract_id_db = get_contract(contract_id)
    client = get_client(contract_id_db["client_id"])["name"]
    employee = get_employee(contract_id_db["employee_id"])["name"]
    service_type = get_type_service(contract_id_db["services_type_id"])["service"]
    return render_template('contract.html', contract=contract_id_db, client=client, employee=employee,
                           service=service_type)


@app.route('/new_contract', methods=['GET'])
def new_contract():
    keys = get_contract(1).keys()
    return render_template('new_contract.html', keys=keys)


@app.route('/new_contract', methods=['POST'])
def new_contract_post():
    add_new_contract(request.form['pos_number'], request.form['pos_services_type_id'], request.form['pos_start_price'],
                     request.form['pos_discount'], request.form['pos_client_id'], request.form['pos_employee_id'])
    return redirect(url_for('contracts'))


@app.route('/update_contract/<int:contract_id>', methods=['GET'])
def update_contract(contract_id):
    contract = get_contract(contract_id)
    return render_template('update_contract.html', contract=contract)


@app.route('/update_contract/<int:contract_id>', methods=['POST'])
def update_contract_post(contract_id):
    add_update_contract(contract_id, request.form['pos_number'], request.form['pos_services_type_id'],
                        request.form['pos_start_price'],
                        request.form['pos_discount'], request.form['pos_client_id'], request.form['pos_employee_id'])
    return redirect(url_for('contract', contract_id=contract_id))


@app.route('/clients')
def clients():
    clients_db = get_clients()
    return render_template('clients.html', clients=clients_db)


@app.route('/client/<int:client_id>')
def client(client_id):
    client_id_db = get_client(client_id)
    return render_template('client.html', client=client_id_db)


@app.route('/new_client', methods=['GET'])
def new_client():
    keys = get_client(1).keys()
    return render_template('new_client.html', keys=keys)


@app.route('/new_client', methods=['POST'])
def new_client_post():
    add_new_client(request.form['pos_name'], request.form['pos_email'], request.form['pos_phone_number'])
    return redirect(url_for('clients'))


@app.route('/update_client/<int:client_id>', methods=['GET'])
def update_client(client_id):
    client = get_client(client_id)
    return render_template('update_client.html', client=client)


@app.route('/update_client/<int:client_id>', methods=['POST'])
def update_client_post(client_id):
    add_update_client(client_id, request.form['pos_name'], request.form['pos_email'],
                      request.form['pos_phone_number'])
    return redirect(url_for('client', client_id=client_id))


@app.route('/employees')
def employees():
    employees_db = get_employees()
    return render_template('employees.html', employees=employees_db)


@app.route('/employee/<int:employee_id>')
def employee(employee_id):
    employee_id_db = get_employee(employee_id)
    return render_template('employee.html', employee=employee_id_db)


@app.route('/new_employee', methods=['GET'])
def new_employee():
    keys = get_employee(1).keys()
    return render_template('new_employee.html', keys=keys)


@app.route('/new_employee', methods=['POST'])
def new_employee_post():
    add_new_employee(request.form['pos_name'], request.form['pos_email'], request.form['pos_phone_number'],
                     request.form['pos_position'])
    return redirect(url_for('employees'))


@app.route('/update_employee/<int:employee_id>', methods=['GET'])
def update_employee(employee_id):
    employ = get_employee(employee_id)
    return render_template('update_employee.html', employee=employ)


@app.route('/update_employee/<int:employee_id>', methods=['POST'])
def update_employee_post(employee_id):
    add_update_employee(employee_id, request.form['pos_name'], request.form['pos_email'],
                        request.form['pos_phone_number'], request.form['pos_position'])
    return redirect(url_for('employee', employee_id=employee_id))


@app.route('/type_services')
def type_services():
    type_services_db = get_type_services()
    return render_template('type_services.html', type_services=type_services_db)


@app.route('/type_service/<int:type_service_id>')
def type_service(type_service_id):
    type_service_id_db = get_type_service(type_service_id)
    return render_template('type_service.html', type_service=type_service_id_db)


@app.route('/new_type_service', methods=['GET'])
def new_type_service():
    keys = get_type_service(1).keys()
    return render_template('new_type_service.html', keys=keys)


@app.route('/new_type_service', methods=['POST'])
def new_type_service_post():
    add_new_type_service(request.form['pos_service'])
    return redirect(url_for('type_services'))


@app.route('/update_type_service/<int:type_service_id>', methods=['GET'])
def update_type_service(type_service_id):
    service = get_type_service(type_service_id)
    return render_template('update_type_service.html', type_service=service)


@app.route('/update_type_service/<int:type_service_id>', methods=['POST'])
def update_type_service_post(type_service_id):
    add_update_type_service(type_service_id, request.form['pos_service'])
    return redirect(url_for('type_service', type_service_id=type_service_id))


@app.route('/generate_contract/<int:contract_id>', methods=['GET'])
def generate_contract(contract_id):
    contract = get_contract(contract_id)
    service = get_type_service(contract["services_type_id"])["service"]
    directory = 'contracts'

    number = contract['number']
    dateee = contract['date']
    filename = f"договор №{number} от {dateee}.docx"
    if filename in os.listdir(directory):
        return send_file(directory + '/' + filename, as_attachment=True)
    return render_template('generate_contract.html', contract=contract, type_service=service)


@app.route('/generate_contract/<int:contract_id>', methods=['POST'])
def generate_contract_post(contract_id):
    create_contract(contract_id, request.form['pos_number'], request.form['pos_services_type_id'],
                    request.form['pos_start_price'],
                    request.form['pos_discount'], request.form['pos_client_id'], request.form['pos_employee_id'])
    return redirect(url_for('contract', contract_id=contract_id))


@app.route('/generate_report', methods=['GET'])
def generate_report():
    return render_template('generate_report.html')


@app.route('/generate_report', methods=['POST'])
def generate_report_post():
    a = select_for_report(request.form['client_id'], request.form['date'])
    client = get_client(request.form['client_id'])
    create_report(a, client, request.form['date'])
    directory = 'reports'
    filename = f"Отчет для клиента №{request.form['client_id']} от {request.form['date']}.docx"
    return send_file(directory + '/' + filename, as_attachment=True)


def create_report(list, client, dat):
    template = 'ShablonReport.docx'
    template_doc = Document(template)
    for paragraph in template_doc.paragraphs:
        if "{{ client_name }}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{ client_name }}", str(client["name"]))
        if "{{ client_number }}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{ client_number }}", client["phone_number"])
        if "{{start_data}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{start_data}}", str(dat))
        if "{{finish_data}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{finish_data}}", str(date.today()))

    for j in range(0, len(list)):
        row_cells = template_doc.tables[0].add_row().cells
        row_cells[0].text = str(j + 1)
        row_cells[1].text = str(get_type_service(list[j]["services_type_id"])["service"])
        row_cells[2].text = str(list[j]["date"])
        row_cells[3].text = str(list[j]["finish_price"])
    cl = client["id"]
    result = f"Отчет для клиента №{cl} от {dat}.docx"
    template_doc.save('reports/' + result)


def create_contract(contract_id, number, service_type, start_price, discount, client_id, employee_id):
    template = 'Shablon.docx'
    contract = get_contract(contract_id)
    client = get_client(client_id)
    employee = get_employee(employee_id)
    template_doc = Document(template)
    for table in template_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "{{contracts_number}}" in cell.text:
                    cell.text = cell.text.replace("{{contracts_number}}", number)
                if "{{client_name}}" in cell.text:
                    cell.text = cell.text.replace("{{client_name}}", client["name"])
                if "{{client_number}}" in cell.text:
                    cell.text = cell.text.replace("{{client_number}}", client["phone_number"])
                if "{{client_email}}" in cell.text:
                    cell.text = cell.text.replace("{{client_email}}", client["email"])
                if "{{type_of_service}}" in cell.text:
                    cell.text = cell.text.replace("{{type_of_service}}", service_type)
                if "{{employees_name}}" in cell.text:
                    cell.text = cell.text.replace("{{employees_name}}", employee["name"])
                if "{{position}}" in cell.text:
                    cell.text = cell.text.replace("{{position}}", employee["position"])
                if "{{start_price}}" in cell.text:
                    cell.text = cell.text.replace("{{start_price}}", start_price)
                if "{{discount}}" in cell.text:
                    cell.text = cell.text.replace("{{discount}}", discount)
                if "{{finish_price}}" in cell.text:
                    cell.text = cell.text.replace("{{finish_price}}",
                                                  str(int(int(start_price) * (1 - int(discount) / 100))))
    number = number
    dateee = contract['date']
    result = f'договор №{number} от {dateee}.docx'
    template_doc.save('contracts/' + result)


if __name__ == '__main__':
    app.run()
